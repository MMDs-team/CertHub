from app.models import Template
from app.serializers.template_serializers import TemplateSerializer
from app.utils.google_func import *
import aspose.words as aw

from credentials import SERVICE_ACCOUNT_FILE

from django.core.files.base import ContentFile

from docx import Document

from pdf2image import convert_from_path

import os, pathlib
import shutil
import subprocess

from threading import Timer

from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status

import zipfile

from django.conf import settings
IMAGES_ROOT = 'static/images/templates/images'
FILES_ROOT = 'static/images/templates/files'

@api_view(['POST'])
def template_history(request):
    user = request.user 
    templates = Template.objects.all().filter(owner=user).order_by('-usage')
    serializers = TemplateSerializer(templates, many=True)
     
    return Response(serializers.data)


@api_view(['GET'])
def get_public_templates(request):
    start_index = int(request.GET.get('from', 0))
    count = int(request.GET.get('count', 20))

    templates = Template.objects.all().filter(is_public=True).order_by('-usage')
    templates = templates[start_index:start_index + count]
    serializers = TemplateSerializer(templates, many=True)

    return Response(serializers.data)


@api_view(['GET'])
def get_template(request, pk):
    try:
        temp = Template.objects.get(pk=pk)
        serializer = TemplateSerializer(temp, many=False)
        return Response(serializer.data)
    except:
        message = {'detail': 'No template found with this id'}
        return Response(message, status=status.HTTP_404_NOT_FOUND)


def update_first_page_image(template):
    file = template.file
    try:
        subprocess.run([
            "soffice",
            "--headless",  # Run in headless mode (no GUI)
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(file.path),
            file.path
        ], check=True)
        print(f"Conversion successful!")
    except subprocess.CalledProcessError as e: print(f"Error during conversion: {e}")
    except FileNotFoundError: print("Error: LibreOffice (soffice) is not installed or not in your PATH.")

    pdf_file = os.path.splitext(file.path)[0] + ".pdf"
    images = convert_from_path(pdf_file, dpi=300, first_page=1, last_page=1)
    image_path = os.path.splitext(file.path)[0]
    if images: images[0].save(image_path, "PNG")

    template.image = image_path
    template.save()

    os.remove(pdf_file)


@api_view(['POST'])
def create_template(request):
    pk = request.data.get('pk')
    file = request.data.get('file') 
    user = request.user 

    temp = Template.objects.create(
        owner = user,
        is_public = False,
        is_active = False,
        usage = 1
    )
    file_name = f'{temp.template_id}.docx'
    image_name = f'{temp.template_id}.png'

    if pk is not None:
        org_temp = Template.objects.get(pk=pk)
        temp.file.save(file_name, org_temp.file)
        temp.image.save(image_name, org_temp.image)
        org_temp.usage += 1 
        org_temp.save()

    elif file is not None:
        temp.file.save(file_name, file)
        update_first_page_image(temp)

    else:
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        temp.file.save(file_name, ContentFile(buffer.read()))
        buffer.close()
        temp.image = None

    temp.save()
    serializer = TemplateSerializer(temp, many=False)
    return Response(serializer.data)


@api_view(['DELETE'])
def remove_template(request, pk):
    try:
        temp = Template.objects.get(pk=pk)

        file_name = f'{temp.template_id}.docx'
        image_name = f'{temp.template_id}.png'

        image_path = os.path.join(settings.BASE_DIR, IMAGES_ROOT, image_name)
        file_path = os.path.join(settings.BASE_DIR, FILES_ROOT, file_name)

        pathlib.Path(image_path).unlink(missing_ok=True)
        pathlib.Path(file_path).unlink(missing_ok=False)
        temp.delete()

        message = {'detail': 'Template deleted'}
        return Response(message, status=status.HTTP_204_NO_CONTENT)
    except:
        message = {'detail': 'No template found with this id'}
        return Response(message, status=status.HTTP_404_NOT_FOUND)


PATTERN = '##'
@api_view(['GET'])
def count_blank_fields(request, pk):
    template = Template.objects.all().filter(template_id=pk).first()
    if template is None: return Response({
        'detail': "There is no document with given id."
    }, status=status.HTTP_204_NO_CONTENT)

    try:
        doc = Document(template.file)
        count = 0
        for paragraph in doc.paragraphs:
            count += paragraph.text.count(PATTERN)

        # Debug
        # print(list_files())
        # delete_all_files()
        # print(list_files())

        return Response({
            'blank_fields': count
        }, status=status.HTTP_200_OK)

    except Exception as e:
        return Response({
            'detail': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def make_doc(request, pk):
    template = Template.objects.all().filter(template_id=pk).first()
    if template is None: return Response({
        'detail': "There is no document with given id."
    }, status=status.HTTP_204_NO_CONTENT)

    try:
        doc_id = None
        if template.is_active: doc_id = find_first_file(str(pk))['id']
        else:
            doc_id = upload_file(doc_name=str(pk), file_path=template.file.path)
            set_doc_public_role(doc_id=doc_id, role='writer')

            template.is_active = True
            template.save()

        return Response({
            'doc_id': doc_id
        }, status=status.HTTP_200_OK)

    except Exception as e:
        return Response({
            'detail': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def donwload_doc(request, pk):
    template = Template.objects.all().filter(template_id=pk).first()
    if template is None: return Response({
        'detail': "There is no document with given id."
    }, status=status.HTTP_204_NO_CONTENT)

    try:
        file = template.file
        download_file(file_path=file.path, doc_id=find_first_file(str(pk))['id'])
        update_first_page_image(template)

        return Response({
            'detail': f'Document file of Template with id={pk} have been updated.'
        }, status=status.HTTP_200_OK)

    except Exception as e:
        return Response({
            'detail': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def delete_zip_file(file_path):
    if os.path.exists(file_path) and os.path.isfile(file_path):
        os.remove(file_path)

INTERVAL = 30
timer = Timer(INTERVAL, delete_zip_file, args=('.zip', ))
@api_view(['POST'])
def get_certs(request, pk):
    required_fields = ['matching', 'cols', 'data']
    for field in required_fields:
        if field not in request.data:
            return Response({
                'detail': f'Missing required field: {field}'
            }, status=status.HTTP_400_BAD_REQUEST)

    template = Template.objects.filter(template_id=pk).first()
    if template is None:
        return Response({
            'detail': "No document found with the given ID."
        }, status=status.HTTP_404_NOT_FOUND)

    def decompose_path(path):
        directory, file_name = os.path.split(path)
        base_name, extension = os.path.splitext(file_name)

        return {
            'path': path,
            'directory': directory,
            'file_name': file_name,
            'base_name': base_name,
            'extension': extension
        }

    matching = [elem for elem in request.data.get('matching')]
    cols = [elem for elem in request.data.get('cols')]
    data = [elem for elem in request.data.get('data')]
    matching = [elem for elem in request.data.get('matching')]
    mp = {col_name: idx for idx, col_name in enumerate(cols)}
    
    doc_info = decompose_path(template.file.path)
    certs = []
    for record in data:
        doc = Document(doc_info['path'])
        for col_name in matching:
            idx = mp[col_name]
            curr_val = str(record[idx]) if idx < len(record) else ''

            for paragraph in doc.paragraphs:
                if PATTERN in paragraph.text:
                    paragraph.text = paragraph.text.replace(PATTERN, curr_val, 1)
                    break

        copy_path = os.path.join(f"{doc_info['directory']}", f"{doc_info['base_name']}_copy{doc_info['extension']}")
        doc.save(copy_path)

        cert_path = os.path.join(f"{doc_info['directory']}", f"{len(certs) + 1}.pdf")
        # aw.Document(copy_path).save(cert_path)
        try:
            subprocess.run([
                "soffice",
                "--headless",  # Run in headless mode (no GUI)
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(cert_path),
                copy_path
            ], check=True)
            shutil.move(os.path.join(f"{doc_info['directory']}", f"{doc_info['base_name']}_copy.pdf"), cert_path)
            certs.append(cert_path)
            print(f"Conversion successful! PDF saved in '{cert_path}'.")
        except subprocess.CalledProcessError as e: print(f"Error during conversion: {e}")
        except FileNotFoundError: print("Error: LibreOffice (soffice) is not installed or not in your PATH.")

        os.remove(copy_path)

    zip_file_path = f"{doc_info['directory']}/{pk}_certs.zip"
    with zipfile.ZipFile(zip_file_path, 'w') as zipf:
        for file_path in certs:
            if os.path.exists(file_path): zipf.write(file_path, arcname=os.path.basename(file_path))

    for cert in certs: os.remove(cert)

    doc = find_first_file(str(pk))
    if doc is not None: delete_file(doc['id'])

    template.is_active = False
    template.save()

    global timer
    timer.cancel()
    timer = Timer(INTERVAL, delete_zip_file, args=(zip_file_path, ))
    timer.start()

    return Response({
        'file_name': f"{template.template_id}_certs.zip"
    }, status=status.HTTP_200_OK)